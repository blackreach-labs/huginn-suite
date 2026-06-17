# app/core/report_customizer.py
"""
Report Customizer Engine
Supports custom report templates with branding, drag-and-drop section ordering,
conditional sections, severity filtering, and multi-format output (PDF, HTML, DOCX, Markdown).
Extends the existing advanced_reporting.py module.
"""
import json
import os
import copy
from datetime import datetime
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field, asdict

from PyQt6.QtCore import QObject, pyqtSignal

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Image,
        Table, TableStyle, PageBreak
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# Severity levels ordered from highest to lowest
SEVERITY_LEVELS = ["critical", "high", "medium", "low", "informational"]

# Default sections available for report templates
DEFAULT_SECTIONS = [
    {"id": "executive_summary", "title": "Executive Summary", "enabled": True,
     "conditional": False, "condition_key": None},
    {"id": "methodology", "title": "Methodology", "enabled": True,
     "conditional": False, "condition_key": None},
    {"id": "findings", "title": "Findings", "enabled": True,
     "conditional": False, "condition_key": None},
    {"id": "risk_matrix", "title": "Risk Matrix", "enabled": True,
     "conditional": False, "condition_key": None},
    {"id": "remediation_plan", "title": "Remediation Plan", "enabled": True,
     "conditional": False, "condition_key": None},
    {"id": "appendices", "title": "Appendices", "enabled": True,
     "conditional": False, "condition_key": None},
    {"id": "attack_coverage", "title": "ATT&CK Coverage", "enabled": True,
     "conditional": True, "condition_key": "attack_mappings"},
]

# Default branding configuration
DEFAULT_BRANDING = {
    "logo_path": "",
    "company_name": "",
    "primary_color": "#2c3e50",
    "secondary_color": "#3498db",
    "header_text": "",
    "footer_text": "",
    "cover_page": True,
}


@dataclass
class ReportSection:
    """Represents a single section in a report template."""
    id: str
    title: str
    enabled: bool = True
    conditional: bool = False
    condition_key: Optional[str] = None

    def to_dict(self) -> Dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: Dict) -> "ReportSection":
        return cls(
            id=data["id"],
            title=data["title"],
            enabled=data.get("enabled", True),
            conditional=data.get("conditional", False),
            condition_key=data.get("condition_key"),
        )


@dataclass
class ReportTemplate:
    """Represents a complete report template configuration."""
    name: str
    sections: List[ReportSection] = field(default_factory=list)
    branding: Dict[str, Any] = field(default_factory=lambda: copy.deepcopy(DEFAULT_BRANDING))
    severity_threshold: str = "low"
    created_at: str = ""
    updated_at: str = ""

    def to_dict(self) -> Dict:
        return {
            "name": self.name,
            "sections": [s.to_dict() for s in self.sections],
            "branding": self.branding,
            "severity_threshold": self.severity_threshold,
            "created_at": self.created_at,
            "updated_at": self.updated_at,
        }

    @classmethod
    def from_dict(cls, data: Dict) -> "ReportTemplate":
        sections = [ReportSection.from_dict(s) for s in data.get("sections", [])]
        return cls(
            name=data["name"],
            sections=sections,
            branding=data.get("branding", copy.deepcopy(DEFAULT_BRANDING)),
            severity_threshold=data.get("severity_threshold", "low"),
            created_at=data.get("created_at", ""),
            updated_at=data.get("updated_at", ""),
        )


class ReportCustomizer(QObject):
    """
    Report Customizer Engine.

    Manages report templates with custom branding, section ordering,
    conditional sections, severity filtering, and multi-format output generation.
    """

    template_created = pyqtSignal(str)       # template_name
    template_saved = pyqtSignal(str)         # template_name
    template_deleted = pyqtSignal(str)       # template_name
    report_generated = pyqtSignal(str)       # output_path
    generation_error = pyqtSignal(str)       # error message

    def __init__(self, templates_dir: str = None):
        super().__init__()
        if templates_dir is None:
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            templates_dir = os.path.join(base_dir, "resources", "templates", "report_templates")
        self.templates_dir = templates_dir
        os.makedirs(self.templates_dir, exist_ok=True)

    def create_template(self, name: str, sections: List[Dict] = None,
                        branding: Dict = None,
                        severity_threshold: str = "low") -> ReportTemplate:
        """
        Create a new report template with the given configuration.

        Args:
            name: Template name (used as filename identifier).
            sections: List of section dicts. If None, uses default sections.
            branding: Branding configuration dict. If None, uses defaults.
            severity_threshold: Minimum severity for finding inclusion.

        Returns:
            The created ReportTemplate instance.

        Raises:
            ValueError: If name is empty or severity_threshold is invalid.
        """
        if not name or not name.strip():
            raise ValueError("Template name cannot be empty")

        threshold_lower = severity_threshold.lower()
        if threshold_lower not in SEVERITY_LEVELS:
            raise ValueError(
                f"Invalid severity threshold: '{severity_threshold}'. "
                f"Must be one of: {', '.join(SEVERITY_LEVELS)}"
            )

        if sections is None:
            section_objects = [ReportSection.from_dict(s) for s in DEFAULT_SECTIONS]
        else:
            section_objects = [ReportSection.from_dict(s) for s in sections]

        if branding is None:
            branding = copy.deepcopy(DEFAULT_BRANDING)

        now = datetime.now().isoformat()
        template = ReportTemplate(
            name=name.strip(),
            sections=section_objects,
            branding=branding,
            severity_threshold=threshold_lower,
            created_at=now,
            updated_at=now,
        )

        self.template_created.emit(template.name)
        return template

    def save_template(self, template: ReportTemplate) -> str:
        """
        Save a report template to disk as JSON.

        Args:
            template: The ReportTemplate to save.

        Returns:
            The file path where the template was saved.

        Raises:
            ValueError: If template name is empty.
        """
        if not template.name or not template.name.strip():
            raise ValueError("Template name cannot be empty")

        template.updated_at = datetime.now().isoformat()
        filename = self._template_filename(template.name)
        filepath = os.path.join(self.templates_dir, filename)

        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(template.to_dict(), f, indent=2)

        self.template_saved.emit(template.name)
        return filepath

    def load_template(self, name: str) -> ReportTemplate:
        """
        Load a report template from disk.

        Args:
            name: The template name to load.

        Returns:
            The loaded ReportTemplate instance.

        Raises:
            FileNotFoundError: If the template file doesn't exist.
            ValueError: If the template JSON is malformed.
        """
        filename = self._template_filename(name)
        filepath = os.path.join(self.templates_dir, filename)

        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Template not found: '{name}'")

        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)

        return ReportTemplate.from_dict(data)

    def delete_template(self, name: str) -> bool:
        """
        Delete a report template from disk.

        Args:
            name: The template name to delete.

        Returns:
            True if deleted, False if not found.
        """
        filename = self._template_filename(name)
        filepath = os.path.join(self.templates_dir, filename)

        if os.path.exists(filepath):
            os.remove(filepath)
            self.template_deleted.emit(name)
            return True
        return False

    def list_templates(self) -> List[str]:
        """
        List all saved template names.

        Returns:
            List of template names (without file extension).
        """
        templates = []
        if os.path.exists(self.templates_dir):
            for filename in os.listdir(self.templates_dir):
                if filename.endswith('.json'):
                    templates.append(filename[:-5].replace('_', ' '))
        return sorted(templates)

    def reorder_sections(self, template: ReportTemplate,
                         new_order: List[str]) -> ReportTemplate:
        """
        Reorder template sections by section ID list (drag-and-drop support).

        Args:
            template: The template to reorder.
            new_order: List of section IDs in the desired order.

        Returns:
            The template with reordered sections.

        Raises:
            ValueError: If new_order contains invalid section IDs.
        """
        existing_ids = {s.id for s in template.sections}
        for section_id in new_order:
            if section_id not in existing_ids:
                raise ValueError(f"Unknown section ID: '{section_id}'")

        section_map = {s.id: s for s in template.sections}
        ordered = [section_map[sid] for sid in new_order if sid in section_map]

        # Append any sections not in new_order at the end
        remaining = [s for s in template.sections if s.id not in new_order]
        template.sections = ordered + remaining
        template.updated_at = datetime.now().isoformat()
        return template

    def generate_report(self, template_name: str, engagement_db,
                        output_path: str, output_format: str = "markdown") -> str:
        """
        Generate a report using the specified template and engagement data.

        Args:
            template_name: Name of the template to use.
            engagement_db: EngagementDatabase instance with data.
            output_path: Path for the output file.
            output_format: One of 'pdf', 'html', 'docx', 'markdown'.

        Returns:
            The output file path.

        Raises:
            FileNotFoundError: If template doesn't exist.
            ValueError: If output_format is unsupported.
            ImportError: If required library for format is unavailable.
        """
        format_lower = output_format.lower()
        supported_formats = ["pdf", "html", "docx", "markdown", "md"]
        if format_lower not in supported_formats:
            raise ValueError(
                f"Unsupported output format: '{output_format}'. "
                f"Supported: {', '.join(supported_formats)}"
            )

        template = self.load_template(template_name)

        # Gather report data from engagement database
        report_data = self._gather_report_data(engagement_db, template)

        # Determine which sections to include
        active_sections = self._resolve_sections(template, report_data)

        # Generate output based on format
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)

        try:
            if format_lower == "pdf":
                self._generate_pdf(template, active_sections, report_data, output_path)
            elif format_lower == "html":
                self._generate_html(template, active_sections, report_data, output_path)
            elif format_lower == "docx":
                self._generate_docx(template, active_sections, report_data, output_path)
            elif format_lower in ("markdown", "md"):
                self._generate_markdown(template, active_sections, report_data, output_path)

            self.report_generated.emit(output_path)
            return output_path
        except Exception as e:
            self.generation_error.emit(str(e))
            raise

    def _gather_report_data(self, engagement_db, template: ReportTemplate) -> Dict[str, Any]:
        """Gather all report data from the engagement database."""
        data: Dict[str, Any] = {
            "findings": [],
            "attack_mappings": [],
            "engagement_meta": {},
            "timeline_entries": [],
            "evidence": [],
        }

        # Get findings filtered by severity threshold
        try:
            all_findings = engagement_db.execute_query(
                "SELECT id, title, severity, description, impact, remediation, "
                "cvss_score, cvss_vector, cwe_id, category, status, created_at "
                "FROM findings ORDER BY id"
            )
            data["findings"] = self._filter_by_severity(all_findings, template.severity_threshold)
        except Exception:
            data["findings"] = []

        # Get ATT&CK mappings
        try:
            data["attack_mappings"] = engagement_db.execute_query(
                "SELECT id, finding_id, technique_id, tactic, procedure_description, status "
                "FROM attack_mappings"
            )
        except Exception:
            data["attack_mappings"] = []

        # Get engagement metadata
        try:
            meta_rows = engagement_db.execute_query(
                "SELECT key, value FROM engagement_meta"
            )
            data["engagement_meta"] = {row[0]: row[1] for row in meta_rows}
        except Exception:
            data["engagement_meta"] = {}

        # Get timeline entries
        try:
            data["timeline_entries"] = engagement_db.execute_query(
                "SELECT id, action_type, actor, description, timestamp "
                "FROM timeline_entries ORDER BY timestamp"
            )
        except Exception:
            data["timeline_entries"] = []

        # Get evidence
        try:
            data["evidence"] = engagement_db.execute_query(
                "SELECT id, evidence_type, title, source_context, created_at "
                "FROM evidence ORDER BY created_at"
            )
        except Exception:
            data["evidence"] = []

        return data

    def _filter_by_severity(self, findings: List, threshold: str) -> List:
        """Filter findings to include only those at or above the severity threshold."""
        threshold_lower = threshold.lower()
        if threshold_lower not in SEVERITY_LEVELS:
            return findings

        threshold_index = SEVERITY_LEVELS.index(threshold_lower)
        # Include findings with severity index <= threshold_index (higher or equal severity)
        filtered = []
        for finding in findings:
            # finding is a tuple: severity is at index 2
            finding_severity = finding[2].lower() if finding[2] else "informational"
            if finding_severity in SEVERITY_LEVELS:
                finding_index = SEVERITY_LEVELS.index(finding_severity)
                if finding_index <= threshold_index:
                    filtered.append(finding)
            else:
                # Unknown severity - include it
                filtered.append(finding)
        return filtered

    def _resolve_sections(self, template: ReportTemplate,
                          report_data: Dict[str, Any]) -> List[ReportSection]:
        """Determine which sections should appear in the report."""
        active = []
        for section in template.sections:
            if not section.enabled:
                continue

            if section.conditional and section.condition_key:
                # Only include if relevant data exists
                condition_data = report_data.get(section.condition_key, [])
                if not condition_data:
                    continue

            active.append(section)
        return active

    def _generate_markdown(self, template: ReportTemplate,
                           sections: List[ReportSection],
                           report_data: Dict[str, Any],
                           output_path: str) -> None:
        """Generate a Markdown report."""
        lines = []

        # Cover page
        if template.branding.get("cover_page"):
            company = template.branding.get("company_name", "")
            if company:
                lines.append(f"# {company}")
                lines.append("")
            lines.append(f"# Security Assessment Report")
            lines.append("")
            lines.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append("")
            lines.append("---")
            lines.append("")

        # Generate each section
        for section in sections:
            content = self._render_section_markdown(section, report_data, template)
            if content:
                lines.append(content)
                lines.append("")

        # Footer
        footer = template.branding.get("footer_text", "")
        if footer:
            lines.append("---")
            lines.append(f"*{footer}*")
            lines.append("")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))

    def _render_section_markdown(self, section: ReportSection,
                                  report_data: Dict[str, Any],
                                  template: ReportTemplate) -> str:
        """Render a single section as Markdown."""
        lines = []
        lines.append(f"## {section.title}")
        lines.append("")

        if section.id == "executive_summary":
            lines.extend(self._render_executive_summary_md(report_data))
        elif section.id == "methodology":
            lines.extend(self._render_methodology_md(report_data))
        elif section.id == "findings":
            lines.extend(self._render_findings_md(report_data))
        elif section.id == "risk_matrix":
            lines.extend(self._render_risk_matrix_md(report_data))
        elif section.id == "remediation_plan":
            lines.extend(self._render_remediation_md(report_data))
        elif section.id == "appendices":
            lines.extend(self._render_appendices_md(report_data))
        elif section.id == "attack_coverage":
            lines.extend(self._render_attack_coverage_md(report_data))

        return "\n".join(lines)

    def _render_executive_summary_md(self, data: Dict) -> List[str]:
        """Render executive summary section in Markdown."""
        lines = []
        findings = data.get("findings", [])
        total = len(findings)

        severity_counts = self._count_severities(findings)
        lines.append(f"**Total Findings:** {total}")
        lines.append("")
        for sev, count in severity_counts.items():
            if count > 0:
                lines.append(f"- **{sev.capitalize()}:** {count}")
        lines.append("")
        return lines

    def _render_methodology_md(self, data: Dict) -> List[str]:
        """Render methodology section in Markdown."""
        lines = []
        lines.append("The assessment was conducted using industry-standard penetration testing methodologies.")
        lines.append("")
        return lines

    def _render_findings_md(self, data: Dict) -> List[str]:
        """Render findings section in Markdown."""
        lines = []
        findings = data.get("findings", [])
        if not findings:
            lines.append("No findings to report.")
            lines.append("")
            return lines

        for finding in findings:
            # finding tuple: (id, title, severity, description, impact, remediation, cvss_score, cvss_vector, cwe_id, category, status, created_at)
            title = finding[1]
            severity = finding[2]
            description = finding[3] or ""
            impact = finding[4] or ""
            remediation = finding[5] or ""
            cvss = finding[6]

            lines.append(f"### {title}")
            lines.append("")
            lines.append(f"**Severity:** {severity.capitalize() if severity else 'Unknown'}")
            if cvss:
                lines.append(f"**CVSS Score:** {cvss}")
            lines.append("")
            if description:
                lines.append(f"**Description:** {description}")
                lines.append("")
            if impact:
                lines.append(f"**Impact:** {impact}")
                lines.append("")
            if remediation:
                lines.append(f"**Remediation:** {remediation}")
                lines.append("")
            lines.append("---")
            lines.append("")

        return lines

    def _render_risk_matrix_md(self, data: Dict) -> List[str]:
        """Render risk matrix section in Markdown."""
        lines = []
        findings = data.get("findings", [])
        severity_counts = self._count_severities(findings)

        lines.append("| Severity | Count |")
        lines.append("|----------|-------|")
        for sev in SEVERITY_LEVELS:
            count = severity_counts.get(sev, 0)
            lines.append(f"| {sev.capitalize()} | {count} |")
        lines.append("")
        return lines

    def _render_remediation_md(self, data: Dict) -> List[str]:
        """Render remediation plan section in Markdown."""
        lines = []
        findings = data.get("findings", [])
        remediation_items = [(f[1], f[2], f[5]) for f in findings if f[5]]

        if not remediation_items:
            lines.append("No remediation items to report.")
            lines.append("")
            return lines

        lines.append("| # | Finding | Severity | Remediation |")
        lines.append("|---|---------|----------|-------------|")
        for i, (title, severity, remediation) in enumerate(remediation_items, 1):
            # Truncate long remediations for table display
            rem_short = remediation[:100] + "..." if len(remediation) > 100 else remediation
            lines.append(f"| {i} | {title} | {severity.capitalize() if severity else 'N/A'} | {rem_short} |")
        lines.append("")
        return lines

    def _render_appendices_md(self, data: Dict) -> List[str]:
        """Render appendices section in Markdown."""
        lines = []
        timeline = data.get("timeline_entries", [])
        if timeline:
            lines.append("### Activity Timeline")
            lines.append("")
            for entry in timeline[:50]:  # Limit to 50 entries
                lines.append(f"- **{entry[4]}** - {entry[3]}")
            lines.append("")
        else:
            lines.append("No appendix data available.")
            lines.append("")
        return lines

    def _render_attack_coverage_md(self, data: Dict) -> List[str]:
        """Render ATT&CK coverage section in Markdown."""
        lines = []
        mappings = data.get("attack_mappings", [])
        if not mappings:
            lines.append("No ATT&CK mappings recorded.")
            lines.append("")
            return lines

        lines.append("| Technique ID | Tactic | Status |")
        lines.append("|-------------|--------|--------|")
        for mapping in mappings:
            # mapping tuple: (id, finding_id, technique_id, tactic, procedure_description, status)
            lines.append(f"| {mapping[2]} | {mapping[3]} | {mapping[5]} |")
        lines.append("")
        return lines

    def _generate_html(self, template: ReportTemplate,
                       sections: List[ReportSection],
                       report_data: Dict[str, Any],
                       output_path: str) -> None:
        """Generate an HTML report with embedded CSS."""
        primary_color = template.branding.get("primary_color", "#2c3e50")
        secondary_color = template.branding.get("secondary_color", "#3498db")
        company_name = template.branding.get("company_name", "")
        header_text = template.branding.get("header_text", "")
        footer_text = template.branding.get("footer_text", "")

        html_parts = []
        html_parts.append(f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Security Assessment Report</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            margin: 0; padding: 0;
            color: #333; line-height: 1.6;
        }}
        .header {{
            background-color: {primary_color};
            color: white; padding: 20px 40px;
            border-bottom: 4px solid {secondary_color};
        }}
        .header h1 {{ margin: 0; font-size: 1.8em; }}
        .header .subtitle {{ opacity: 0.8; margin-top: 5px; }}
        .content {{ padding: 40px; max-width: 1000px; margin: 0 auto; }}
        h2 {{
            color: {primary_color};
            border-bottom: 2px solid {secondary_color};
            padding-bottom: 8px;
            margin-top: 40px;
        }}
        h3 {{ color: {primary_color}; }}
        table {{
            width: 100%; border-collapse: collapse;
            margin: 15px 0;
        }}
        th, td {{
            border: 1px solid #ddd; padding: 10px;
            text-align: left;
        }}
        th {{ background-color: {primary_color}; color: white; }}
        .finding {{
            border-left: 4px solid {secondary_color};
            padding: 15px; margin: 15px 0;
            background-color: #f9f9f9;
        }}
        .severity-critical {{ border-left-color: #e74c3c; }}
        .severity-high {{ border-left-color: #e67e22; }}
        .severity-medium {{ border-left-color: #f39c12; }}
        .severity-low {{ border-left-color: #27ae60; }}
        .footer {{
            background-color: #ecf0f1;
            padding: 15px 40px; text-align: center;
            font-size: 0.9em; color: #666;
            border-top: 2px solid {secondary_color};
        }}
        .cover-page {{
            text-align: center; padding: 100px 40px;
            background-color: {primary_color}; color: white;
            min-height: 400px;
            display: flex; flex-direction: column;
            justify-content: center; align-items: center;
        }}
        .cover-page h1 {{ font-size: 2.5em; margin-bottom: 20px; }}
        .cover-page .date {{ opacity: 0.8; font-size: 1.2em; }}
    </style>
</head>
<body>
""")

        # Cover page
        if template.branding.get("cover_page"):
            html_parts.append(f"""<div class="cover-page">
    <h1>{company_name or 'Security Assessment Report'}</h1>
    <p class="date">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
</div>
""")

        # Header
        if header_text:
            html_parts.append(f"""<div class="header">
    <h1>{header_text}</h1>
</div>
""")

        html_parts.append('<div class="content">')

        # Render sections
        for section in sections:
            html_parts.append(self._render_section_html(section, report_data, template))

        html_parts.append('</div>')

        # Footer
        if footer_text:
            html_parts.append(f'<div class="footer"><p>{footer_text}</p></div>')

        html_parts.append('</body>\n</html>')

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(html_parts))

    def _render_section_html(self, section: ReportSection,
                              report_data: Dict[str, Any],
                              template: ReportTemplate) -> str:
        """Render a single section as HTML."""
        parts = [f'<h2>{section.title}</h2>']

        if section.id == "executive_summary":
            parts.extend(self._render_executive_summary_html(report_data))
        elif section.id == "findings":
            parts.extend(self._render_findings_html(report_data))
        elif section.id == "risk_matrix":
            parts.extend(self._render_risk_matrix_html(report_data))
        elif section.id == "methodology":
            parts.append("<p>The assessment was conducted using industry-standard penetration testing methodologies.</p>")
        elif section.id == "remediation_plan":
            parts.extend(self._render_remediation_html(report_data))
        elif section.id == "appendices":
            parts.extend(self._render_appendices_html(report_data))
        elif section.id == "attack_coverage":
            parts.extend(self._render_attack_coverage_html(report_data))

        return "\n".join(parts)

    def _render_executive_summary_html(self, data: Dict) -> List[str]:
        """Render executive summary section in HTML."""
        findings = data.get("findings", [])
        severity_counts = self._count_severities(findings)
        parts = [f"<p><strong>Total Findings:</strong> {len(findings)}</p>", "<ul>"]
        for sev, count in severity_counts.items():
            if count > 0:
                parts.append(f"<li><strong>{sev.capitalize()}:</strong> {count}</li>")
        parts.append("</ul>")
        return parts

    def _render_findings_html(self, data: Dict) -> List[str]:
        """Render findings section in HTML."""
        findings = data.get("findings", [])
        if not findings:
            return ["<p>No findings to report.</p>"]

        parts = []
        for finding in findings:
            severity = (finding[2] or "unknown").lower()
            parts.append(f'<div class="finding severity-{severity}">')
            parts.append(f"<h3>{finding[1]}</h3>")
            parts.append(f"<p><strong>Severity:</strong> {severity.capitalize()}</p>")
            if finding[6]:
                parts.append(f"<p><strong>CVSS Score:</strong> {finding[6]}</p>")
            if finding[3]:
                parts.append(f"<p><strong>Description:</strong> {finding[3]}</p>")
            if finding[4]:
                parts.append(f"<p><strong>Impact:</strong> {finding[4]}</p>")
            if finding[5]:
                parts.append(f"<p><strong>Remediation:</strong> {finding[5]}</p>")
            parts.append("</div>")
        return parts

    def _render_risk_matrix_html(self, data: Dict) -> List[str]:
        """Render risk matrix as HTML table."""
        findings = data.get("findings", [])
        severity_counts = self._count_severities(findings)
        parts = ["<table>", "<tr><th>Severity</th><th>Count</th></tr>"]
        for sev in SEVERITY_LEVELS:
            count = severity_counts.get(sev, 0)
            parts.append(f"<tr><td>{sev.capitalize()}</td><td>{count}</td></tr>")
        parts.append("</table>")
        return parts

    def _render_remediation_html(self, data: Dict) -> List[str]:
        """Render remediation plan as HTML table."""
        findings = data.get("findings", [])
        items = [(f[1], f[2], f[5]) for f in findings if f[5]]
        if not items:
            return ["<p>No remediation items to report.</p>"]

        parts = ["<table>", "<tr><th>#</th><th>Finding</th><th>Severity</th><th>Remediation</th></tr>"]
        for i, (title, severity, remediation) in enumerate(items, 1):
            parts.append(f"<tr><td>{i}</td><td>{title}</td><td>{severity.capitalize() if severity else 'N/A'}</td><td>{remediation}</td></tr>")
        parts.append("</table>")
        return parts

    def _render_appendices_html(self, data: Dict) -> List[str]:
        """Render appendices section in HTML."""
        timeline = data.get("timeline_entries", [])
        if not timeline:
            return ["<p>No appendix data available.</p>"]
        parts = ["<h3>Activity Timeline</h3>", "<ul>"]
        for entry in timeline[:50]:
            parts.append(f"<li><strong>{entry[4]}</strong> - {entry[3]}</li>")
        parts.append("</ul>")
        return parts

    def _render_attack_coverage_html(self, data: Dict) -> List[str]:
        """Render ATT&CK coverage section in HTML."""
        mappings = data.get("attack_mappings", [])
        if not mappings:
            return ["<p>No ATT&CK mappings recorded.</p>"]
        parts = ["<table>", "<tr><th>Technique ID</th><th>Tactic</th><th>Status</th></tr>"]
        for mapping in mappings:
            parts.append(f"<tr><td>{mapping[2]}</td><td>{mapping[3]}</td><td>{mapping[5]}</td></tr>")
        parts.append("</table>")
        return parts

    def _generate_pdf(self, template: ReportTemplate,
                      sections: List[ReportSection],
                      report_data: Dict[str, Any],
                      output_path: str) -> None:
        """Generate a PDF report using reportlab."""
        if not REPORTLAB_AVAILABLE:
            raise ImportError(
                "reportlab is required for PDF generation. "
                "Install with: pip install reportlab"
            )

        doc = SimpleDocTemplate(output_path, pagesize=letter,
                                leftMargin=inch, rightMargin=inch,
                                topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()
        story = []

        primary_color = template.branding.get("primary_color", "#2c3e50")
        company_name = template.branding.get("company_name", "")

        # Custom styles
        title_style = ParagraphStyle(
            'ReportTitle', parent=styles['Heading1'],
            fontSize=24, spaceAfter=30,
            textColor=HexColor(primary_color)
        )
        heading_style = ParagraphStyle(
            'ReportHeading', parent=styles['Heading2'],
            fontSize=16, spaceAfter=12,
            textColor=HexColor(primary_color)
        )

        # Cover page
        if template.branding.get("cover_page"):
            logo_path = template.branding.get("logo_path", "")
            if logo_path and os.path.exists(logo_path):
                try:
                    story.append(Image(logo_path, width=2*inch, height=2*inch))
                    story.append(Spacer(1, 20))
                except Exception:
                    pass

            if company_name:
                story.append(Paragraph(company_name, title_style))

            story.append(Paragraph("Security Assessment Report", title_style))
            story.append(Spacer(1, 20))
            story.append(Paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                styles['Normal']
            ))
            story.append(PageBreak())

        # Generate sections
        for section in sections:
            story.append(Paragraph(section.title, heading_style))
            story.append(Spacer(1, 12))

            section_content = self._render_section_pdf(section, report_data, styles)
            story.extend(section_content)
            story.append(Spacer(1, 20))

        doc.build(story)

    def _render_section_pdf(self, section: ReportSection,
                             report_data: Dict[str, Any],
                             styles) -> List:
        """Render a section for PDF output."""
        elements = []

        if section.id == "executive_summary":
            findings = report_data.get("findings", [])
            severity_counts = self._count_severities(findings)
            elements.append(Paragraph(f"Total Findings: {len(findings)}", styles['Normal']))
            for sev, count in severity_counts.items():
                if count > 0:
                    elements.append(Paragraph(f"  • {sev.capitalize()}: {count}", styles['Normal']))

        elif section.id == "findings":
            findings = report_data.get("findings", [])
            if not findings:
                elements.append(Paragraph("No findings to report.", styles['Normal']))
            else:
                for finding in findings:
                    elements.append(Paragraph(f"<b>{finding[1]}</b>", styles['Heading3']))
                    elements.append(Paragraph(
                        f"Severity: {finding[2].capitalize() if finding[2] else 'Unknown'}",
                        styles['Normal']
                    ))
                    if finding[3]:
                        elements.append(Paragraph(f"Description: {finding[3]}", styles['Normal']))
                    elements.append(Spacer(1, 8))

        elif section.id == "risk_matrix":
            findings = report_data.get("findings", [])
            severity_counts = self._count_severities(findings)
            table_data = [["Severity", "Count"]]
            for sev in SEVERITY_LEVELS:
                table_data.append([sev.capitalize(), str(severity_counts.get(sev, 0))])
            t = Table(table_data, colWidths=[3*inch, 2*inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), HexColor("#2c3e50")),
                ('TEXTCOLOR', (0, 0), (-1, 0), HexColor("#ffffff")),
                ('GRID', (0, 0), (-1, -1), 1, HexColor("#dddddd")),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
            ]))
            elements.append(t)

        elif section.id == "methodology":
            elements.append(Paragraph(
                "The assessment was conducted using industry-standard penetration testing methodologies.",
                styles['Normal']
            ))

        elif section.id == "remediation_plan":
            findings = report_data.get("findings", [])
            items = [(f[1], f[2], f[5]) for f in findings if f[5]]
            if not items:
                elements.append(Paragraph("No remediation items to report.", styles['Normal']))
            else:
                for i, (title, severity, remediation) in enumerate(items, 1):
                    elements.append(Paragraph(
                        f"{i}. <b>{title}</b> ({severity.capitalize() if severity else 'N/A'}): {remediation}",
                        styles['Normal']
                    ))

        elif section.id == "attack_coverage":
            mappings = report_data.get("attack_mappings", [])
            if not mappings:
                elements.append(Paragraph("No ATT&CK mappings recorded.", styles['Normal']))
            else:
                table_data = [["Technique ID", "Tactic", "Status"]]
                for m in mappings:
                    table_data.append([m[2], m[3], m[5]])
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), HexColor("#2c3e50")),
                    ('TEXTCOLOR', (0, 0), (-1, 0), HexColor("#ffffff")),
                    ('GRID', (0, 0), (-1, -1), 1, HexColor("#dddddd")),
                ]))
                elements.append(t)

        elif section.id == "appendices":
            timeline = report_data.get("timeline_entries", [])
            if timeline:
                for entry in timeline[:50]:
                    elements.append(Paragraph(f"• {entry[4]} - {entry[3]}", styles['Normal']))
            else:
                elements.append(Paragraph("No appendix data available.", styles['Normal']))

        return elements

    def _generate_docx(self, template: ReportTemplate,
                       sections: List[ReportSection],
                       report_data: Dict[str, Any],
                       output_path: str) -> None:
        """Generate a DOCX report using python-docx."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install with: pip install python-docx"
            )

        doc = DocxDocument()
        company_name = template.branding.get("company_name", "")

        # Cover page
        if template.branding.get("cover_page"):
            if company_name:
                title = doc.add_heading(company_name, level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            title = doc.add_heading("Security Assessment Report", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

        # Generate sections
        for section in sections:
            doc.add_heading(section.title, level=1)
            self._render_section_docx(doc, section, report_data)

        doc.save(output_path)

    def _render_section_docx(self, doc, section: ReportSection,
                              report_data: Dict[str, Any]) -> None:
        """Render a section in DOCX format."""
        if section.id == "executive_summary":
            findings = report_data.get("findings", [])
            severity_counts = self._count_severities(findings)
            doc.add_paragraph(f"Total Findings: {len(findings)}")
            for sev, count in severity_counts.items():
                if count > 0:
                    doc.add_paragraph(f"{sev.capitalize()}: {count}", style='List Bullet')

        elif section.id == "findings":
            findings = report_data.get("findings", [])
            if not findings:
                doc.add_paragraph("No findings to report.")
            else:
                for finding in findings:
                    doc.add_heading(finding[1], level=2)
                    doc.add_paragraph(f"Severity: {finding[2].capitalize() if finding[2] else 'Unknown'}")
                    if finding[3]:
                        doc.add_paragraph(f"Description: {finding[3]}")
                    if finding[5]:
                        doc.add_paragraph(f"Remediation: {finding[5]}")

        elif section.id == "methodology":
            doc.add_paragraph(
                "The assessment was conducted using industry-standard penetration testing methodologies."
            )

        elif section.id == "risk_matrix":
            findings = report_data.get("findings", [])
            severity_counts = self._count_severities(findings)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Severity"
            hdr[1].text = "Count"
            for sev in SEVERITY_LEVELS:
                row = table.add_row().cells
                row[0].text = sev.capitalize()
                row[1].text = str(severity_counts.get(sev, 0))

        elif section.id == "remediation_plan":
            findings = report_data.get("findings", [])
            items = [(f[1], f[2], f[5]) for f in findings if f[5]]
            if not items:
                doc.add_paragraph("No remediation items to report.")
            else:
                for i, (title, severity, remediation) in enumerate(items, 1):
                    doc.add_paragraph(
                        f"{i}. {title} ({severity.capitalize() if severity else 'N/A'}): {remediation}"
                    )

        elif section.id == "attack_coverage":
            mappings = report_data.get("attack_mappings", [])
            if not mappings:
                doc.add_paragraph("No ATT&CK mappings recorded.")
            else:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = "Technique ID"
                hdr[1].text = "Tactic"
                hdr[2].text = "Status"
                for m in mappings:
                    row = table.add_row().cells
                    row[0].text = m[2]
                    row[1].text = m[3]
                    row[2].text = m[5]

        elif section.id == "appendices":
            timeline = report_data.get("timeline_entries", [])
            if timeline:
                doc.add_heading("Activity Timeline", level=2)
                for entry in timeline[:50]:
                    doc.add_paragraph(f"{entry[4]} - {entry[3]}", style='List Bullet')
            else:
                doc.add_paragraph("No appendix data available.")

    def _count_severities(self, findings: List) -> Dict[str, int]:
        """Count findings by severity level."""
        counts = {sev: 0 for sev in SEVERITY_LEVELS}
        for finding in findings:
            severity = (finding[2] or "informational").lower()
            if severity in counts:
                counts[severity] += 1
        return counts

    def _template_filename(self, name: str) -> str:
        """Convert template name to safe filename."""
        safe_name = name.strip().replace(' ', '_').lower()
        # Remove unsafe characters
        safe_name = ''.join(c for c in safe_name if c.isalnum() or c in ('_', '-'))
        return f"{safe_name}.json"
